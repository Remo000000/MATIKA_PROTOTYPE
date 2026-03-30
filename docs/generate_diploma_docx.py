# -*- coding: utf-8 -*-
"""Build diploma .docx from diploma_matika.md + university BЖ formatting (TNR 14, margins)."""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Mm, Pt

BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "diploma_matika.md"
OUT_PATH = BASE / "Дипломка документация.docx"

# Ереже 8.5.1.2: аңдатпа ≤150 таңба; 8.5.1.6: қорытынды ≤700 таңба
ANNOT_KK = (
    "MATIKA (Django): оқу кестесін оңтайландыру; слотқа нейро болжам, greedy/GA; SQLite немесе PostgreSQL. "
    "Мақсат — деректер мен ML кестеге енгізілуі."
)
assert len(ANNOT_KK) <= 150

CONCLUSION_KK = (
    "MATIKA веб-жүйесінің архитектурасы мен іске асырылуы сипатталды: Django, рөлдер, тенант, "
    "кесте генерациясы мен GA, слоттың қолайсыздығын Keras немесе эвристика арқылы есептеу. "
    "Прототип сценарийлер бойынша тексерілді; шектеулер мен даму бағыттары көрсетілді. "
    "Нәтиже ЖОО-да пилот немесе кеңейту үшін негіз бола алады; нақты LMS интеграциясы мен "
    "өндірістік талаптар келесі кезеңде қарастырылады."
)
assert len(CONCLUSION_KK) <= 700


def set_doc_defaults(doc: Document) -> None:
    sec = doc.sections[0]
    sec.left_margin = Mm(30)
    sec.right_margin = Mm(10)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Mm(12.5)


def heading_center(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def _set_cell_run(cell, text: str, *, bold: bool = False, size_pt: int = 12) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_table_caption(doc: Document, caption: str) -> None:
    """Ереже 8.6.1.11: кесте нөмірі мен қысқа атау сол жақта, абзац шегініссіз."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Mm(0)
    p.paragraph_format.left_indent = Mm(0)
    r = p.add_run(caption)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)


def add_data_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]]) -> None:
    add_table_caption(doc, caption)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_run(hdr[i], h, bold=True, size_pt=12)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _set_cell_run(tbl.rows[ri + 1].cells[ci], val, bold=False, size_pt=12)
    doc.add_paragraph()


def add_diploma_tables(doc: Document) -> None:
    """Тараулар мазмұнына сәйкес 5 кесте: 1.1, 2.1, 2.2, 2.6/2.3, 2.7."""
    add_para(
        doc,
        "Төмендегі кестелер 1-бөлім мен 2-бөлімде сипатталған тәсілдер, архитектура, модульдер, "
        "слот белгілері және генерация шектеулері бойынша мәліметтерді жинақтау үшін берілді.",
    )
    doc.add_paragraph()

    # Кесте 1 — 1.1 тәсілдерді салыстыру
    add_data_table(
        doc,
        "Кесте 1 — Оқу кестесін құру тәсілдерінің салыстырмалы сипаттамасы (1.1-тарау)",
        ["Тәсіл", "Негізгі сипаты", "Тиімділік пен шектеулер", "MATIKA-мен сәйкестігі"],
        [
            [
                "Қолмен немесе Excel",
                "Кесте толтыру, өзгерістерді қолмен енгізу",
                "Уақыт шығыны, қате қаупі; слот сапасын деректермен автоматты бағалау жоқ",
                "Деректер мен генератор арқылы қайталанатын сценарий",
            ],
            [
                "Корпоративтік САЖ класы",
                "Студент/оқытушы деректерін орталықтандыру",
                "Интеграция тереңдігі әртүрлі; слотқа нейро болжам әрқашан болмайды",
                "Тенант + анықтамалар; слот болжамын қосу",
            ],
            [
                "AI/кесте генераторлары",
                "Шектеулер + деректер/эвристика",
                "Вузбен бейімдеу және локалды ерекшеліктер қажет",
                "Greedy + GA + Keras/эвристика",
            ],
        ],
    )

    # Кесте 2 — 2.1 үш қабат
    add_data_table(
        doc,
        "Кесте 2 — MATIKA клиент-сервер архитектурасының логикалық қабаттары (2.1-тарау)",
        ["Қабат", "Негізгі құрам", "Мақсаты"],
        [
            [
                "Көрсету",
                "Django шаблондары, static, Bootstrap, Chart.js",
                "Пайдаланушы интерфейсі, навигация, кесте көрінісі",
            ],
            [
                "Қолданба логикасы",
                "views, forms, services, ML инференс, REST API",
                "Рөлдер, генерация, экспорт, слот болжамы",
            ],
            [
                "Деректер",
                "Django ORM, SQLite / PostgreSQL",
                "Анықтамалар, сабақ жазбалары, слот белгілері, журналдар",
            ],
        ],
    )

    # Кесте 3 — 2.2 модульдер
    add_data_table(
        doc,
        "Кесте 3 — Django қолданбаларының функциялық міндеттері (2.2-тарау)",
        ["Қолданба", "Негізгі функциялар", "Бағытталған URL / аймақ"],
        [
            [
                "accounts",
                "Кіру, тіркелу, хабарламалар, профиль өзгерістерін келісу, әрекет журналы",
                "/accounts/",
            ],
            [
                "university",
                "Ұйым (тенант), факультет, кафедра, топ, аудитория, слот, кезең; CSV импорт",
                "/university/",
            ],
            [
                "scheduling",
                "Оқу талаптары, сабақтар, жеке кесте, генерация, GA, тілектер, экспорт, API",
                "/scheduling/",
            ],
            [
                "scheduling.ml",
                "Слоттың ыңғайсыздығын болжау (Keras немесе эвристика), CSV",
                "/scheduling/slot-prediction/",
            ],
            [
                "dashboard",
                "Басты бет, аналитика, CSV экспорт",
                "/analytics/ және басты маршруттар",
            ],
        ],
    )

    # Кесте 4 — слот белгілері 2.6 / SlotPedagogicalFeatures
    add_data_table(
        doc,
        "Кесте 4 — Слоттың педагогикалық белгілері мен деректер көздері (2.3, 2.6-тараулар)",
        ["Көрсеткіш / белгі", "Мазмұны", "Ескерту"],
        [
            ["Апта күні, сабақ нөмірі (period)", "Уақыт торындағы позиция", "Нормалдау, кірісте қолданылады"],
            ["Шаршаңдылық, сауалнама жүктемесі", "Субъективті жүктеме көрсеткіштері", "Сауалнама / әкімші енгізімі"],
            ["LMS белсенділігі", "Онлайн белсенділік индикаторы", "LMS интеграциясы перспективасы"],
            ["Тарихи семестр жүктемесі", "Өткен кезеңдер статистикасы", "Дерекқордағы SlotPedagogicalFeatures"],
            ["Мақсатты белгі (болжам үшін)", "Модельді оқыту нысаны", "Нейро болжам қосылған жағдайда"],
        ],
    )

    # Кесте 5 — қатал/жұмсақ 2.7
    add_data_table(
        doc,
        "Кесте 5 — Кесте генерациясының қатал шектеулері мен жұмсақ айыптар (2.7, 2.6-тараулар)",
        ["Түрі", "Мысал шарттары", "Генератордағы рөлі"],
        [
            [
                "Қатал",
                "Бір слотта бір оқытушы; бір топқа бір сабақ; бір аудиторияға бір сабақ",
                "Бұзылмас шарт, орналастыру мүмкін емес",
            ],
            [
                "Қатал",
                "Аудитория сыйымдылығы ≥ топ және TeachingRequirement минимумы",
                "Үміткер аудиторияларды сүзу",
            ],
            [
                "Жұмсақ",
                "Оқытушы күні/сағатына сәйкестік, «терезелер», қатардағы сабақтар",
                "Айып функциясы, GA фитнес",
            ],
            [
                "Жұмсақ",
                "Слоттың ыңғайсыздығы 0…1 (Keras немесе эвристика)",
                "ml_penalty_units, greedy реттеу, GA",
            ],
            [
                "Қызметтік",
                "Мұздатылған (is_frozen) сабақтар",
                "GA/жергілікті жақсартуда өзгертілмейді",
            ],
        ],
    )


def strip_md_inline(s: str) -> str:
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def parse_md_body(md: str) -> str:
    """Кіріспеден Қорытындыға дейін (қоспай)."""
    start = md.find("## Кіріспе")
    end = md.find("## Қорытынды")
    if start == -1 or end == -1:
        raise ValueError("MD: Кіріспе/Қорытынды бөлімдері табылмады")
    return md[start:end].strip()


def md_blocks_to_doc(doc: Document, body: str) -> None:
    """## / ### және абзацтарды Word-қа шығарады."""
    lines = body.splitlines()
    i = 0
    first_h2 = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("## ") and not line.startswith("###"):
            title = line[3:].strip()
            if not first_h2:
                page_break(doc)
            first_h2 = False
            heading_center(doc, title.upper())
            i += 1
            continue
        if line.startswith("### "):
            add_para(doc, line[4:].strip(), bold=True)
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        # жинақта абзац (келесі ### немесе ## дейін)
        buf = [strip_md_inline(line)]
        i += 1
        while i < len(lines):
            n = lines[i].rstrip()
            if not n:
                i += 1
                break
            if n.startswith("#"):
                break
            buf.append(strip_md_inline(n))
            i += 1
        text = " ".join(buf).replace("  ", " ").strip()
        if text:
            add_para(doc, text)


def parse_references(md: str) -> list[str]:
    sec = md.find("## Пайдаланған әдебиеттер")
    if sec == -1:
        return []
    rest = md[sec:].split("---", 1)[0]
    out = []
    for line in rest.splitlines():
        m = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if m:
            t = m.group(2).strip()
            if t.startswith("(") and "PDF" in t:
                continue
            out.append(t)
    # толықтырулар (жартылай атаулар)
    fix = {
        "Predicting the Performance of Students Using Deep.": "Predicting the Performance of Students Using Deep Learning.",
        "Enhancing Student Performance Prediction on Learnersourced Questions with.": (
            "Enhancing student performance prediction on learner-sourced questions with deep learning."
        ),
        "Graph Neural Network Heuristic for.": (
            "Graph neural network heuristic for the construction of initial solutions to educational timetabling problems."
        ),
        "Using Deep Learning in Student Performance.": "Using deep learning in student performance prediction.",
    }
    return [fix.get(x, x) for x in out]


def main() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    body = parse_md_body(md)
    refs = parse_references(md)

    doc = Document()
    set_doc_defaults(doc)

    heading_center(doc, "«ҚАЗАҚ ҰЛТТЫҚ ҚЫЗДАР ПЕДАГОГИКАЛЫҚ УНИВЕРСИТЕТІ» КеАҚ")
    doc.add_paragraph()
    heading_center(doc, "БІТІРУ ЖОБАСЫ")
    doc.add_paragraph()
    add_para(
        doc,
        "Тақырыбы: нейрондық желілермен деректерді талдау негізінде сабақтардың оңтайлы уақыт "
        "аралықтарын болжай отырып, оқу жоспарлары мен кестелерді басқаруға арналған веб-жүйені "
        "әзірлеу (MATIKA)",
    )
    doc.add_paragraph()
    add_para(doc, "Орындаушы: _________________________________")
    add_para(doc, "Ғылыми жетекші: _________________________________")
    add_para(doc, "Кафедра / білім беру бағдарламасы: _________________________________")
    add_para(doc, "Алматы, 2026 ж.")
    page_break(doc)

    heading_center(doc, "АҢДАТПА")
    add_para(doc, ANNOT_KK)
    page_break(doc)

    heading_center(doc, "АННОТАЦИЯ (орысша)")
    ru = re.search(
        r"## АННОТАЦИЯ.*?\n\n(.+?)\n\n---",
        md,
        re.DOTALL,
    )
    if ru:
        add_para(doc, ru.group(1).strip().replace("\n", " "))
    page_break(doc)

    heading_center(doc, "ANNOTATION (English)")
    en = re.search(
        r"## ANNOTATION.*?\n\n(.+?)\n\n---",
        md,
        re.DOTALL,
    )
    if en:
        add_para(doc, en.group(1).strip().replace("\n", " "))
    page_break(doc)

    heading_center(doc, "МАЗМҰНЫ")
    for line in (
        "Кіріспе ..........................................................................",
        "1 Пәндік саланы талдау ......................................................",
        "2 Жобалау және әзірлеу ........................................................",
        "Кестелер жинағы ..............................................................",
        "Қорытынды ......................................................................",
        "Пайдаланған әдебиеттер ......................................................",
        "Қосымшалар .....................................................................",
    ):
        add_para(doc, line)
    add_para(
        doc,
        "Ескерту: Word-та «Мазмұн» өрісін қолданып бет нөмірлерін жаңартыңыз (Ереже 8.5.1.3).",
    )
    page_break(doc)

    # Негізгі мәтін: кіріспе + 1 + 2 бөлім (толық md)
    md_blocks_to_doc(doc, body)

    page_break(doc)
    heading_center(doc, "КЕСТЕЛЕР ЖИНАҒЫ")
    add_diploma_tables(doc)

    page_break(doc)
    heading_center(doc, "ҚОРЫТЫНДЫ")
    add_para(doc, CONCLUSION_KK)
    page_break(doc)

    heading_center(doc, "ПАЙДАЛАНҒАН ӘДЕБИЕТТЕР")
    add_para(
        doc,
        "Тізім реттік нөмірмен рәсімделеді (Ереже 8.5.1.7). APA толық сипаты үшін 3-қосымшаны "
        "қолданып, автор, жыл, басылым, DOI қосыңыз.",
    )
    for i, r in enumerate(refs, start=1):
        add_para(doc, f"{i}. {r}")

    page_break(doc)
    heading_center(doc, "ҚОСЫМША А")
    add_para(doc, "UML: `docs/plantuml/` немесе PNG/SVG экспорт.")
    page_break(doc)
    heading_center(doc, "ҚОСЫМША Б")
    add_para(
        doc,
        "Скриншоттар, кесте мысалы, код үзінділері (кафедра талабына сәйкес).",
    )

    doc.save(OUT_PATH)
    print("OK:", OUT_PATH)


if __name__ == "__main__":
    main()
